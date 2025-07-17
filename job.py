import datetime
import pandas as pd
from docx import Document

output_dir = "output/"
input_file = "volontari.xlsx"
template_file = "scheda.docx"

# Controlla che nell'excel il nome volontario sia nella prima colonna
# e il cognome nella seconda colonna

# nome della colonna excel con la sede di attuazione
nomeColonnaSede = "SEDE"
nomeColonnaProgetto = "PROGETTO"

# testo da inserire nel doc. prima del nome e cognome del volontario
# es: "Nominativo op. vol.: Luigi Nerone"
etichettaVolontario = "Nominativo op. vol.:"

# qui inserire le relazioni, cioè dove andare a prendere la colonna
# e dove mettere il valore:
# testo sul doc. : colonna sul excel
replace_fields = {
    "Data di avvio:": "data avvio",
    "Data di conclusione:": "data fine",
    "Titolo del progetto:": nomeColonnaProgetto,
    "Sede di attuazione:": nomeColonnaSede,
    "OLP - Operatorǝ Locale di Progetto:": "OLP",
    "Telefono OLP:": "TELEFONO OLP",
    "E.mail OLP:": "EMAIL OLP",
    "Docentɜ di formazione specifica del progetto:": "FORMATOR3",
}

check = []


# sostituisce nel file new_document che contiene paragrafi e tabelle
# una testo word con un testo replacement
def replace(new_document, word, replacement):
    # print("replace", word, replacement)
    for p in new_document.paragraphs:
        if p.text.find(word) >= 0:
            p.text = p.text.replace(word, replacement)

    for table in new_document.tables:
        for r in table.rows:
            for c in r.cells:
                if c.text.find(word) >= 0:
                    c.text = c.text.replace(word, replacement)


# Crea il file "sede_nome_cognome.docx" riempiendo gli spazi indicati
def create_file(row):
    nome = row.iloc[0]
    cognome = row.iloc[1]
    sede = (
        str(row.loc[nomeColonnaSede])
        .replace("/", "-")
        .replace(",", "-")
        .replace("\n", "")
        .replace(":", "-")
    )
    progetto = (
        str(row.loc[nomeColonnaProgetto])
        .replace("/", "-")
        .replace(",", "-")
        .replace("\n", "")
        .replace(":", "-")
    )
    filename = progetto + "_" + nome + "_" + cognome + ".docx"

    new_document = Document(template_file)
    replace(
        new_document,
        etichettaVolontario,
        etichettaVolontario + " " + nome + " " + cognome,
    )
    for word in replace_fields:
        originalValue = row.loc[replace_fields[word]]
        if isinstance(originalValue, datetime.datetime):
            originalValue = originalValue.strftime("%d/%m/%Y")
        value = str(originalValue).strip()
        if value == "" or value == "nan":
            check.append(nome + " " + cognome + " manca " + replace_fields[word])
        else:
            replace(new_document, word, word + " " + value)

    new_document.save(output_dir + filename)


## main
df = pd.read_excel(input_file, 0)
# decommenta questo e commenta quello dopo per testare
# df.head(10).apply(create_file, axis=1)
df.apply(create_file, axis=1)

if len(check):
    print("Possibili ERRORI:")
    print("\n".join(check))
