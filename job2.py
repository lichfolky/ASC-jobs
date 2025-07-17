import datetime
import pandas as pd
from docx import Document

"""
per ogni sede fare un file prendendendo 
ogni file ha uno o più progetti indicati nel progetti.xlsx

nel file inserire un elenco con per ogni progetto:
- titolo programma 
- titolo progetto - x posti ordinari e y posti GMO  
"""


output_dir = "output/"
input_file = "progetti.xlsx"
template_file = "sedi.docx"
check = []


def clean(txt):
    return (
        str(txt)
        .strip()
        .replace("/", "-")
        .replace(",", "-")
        .replace(":", "-")
        .replace('"', "")
        .replace("  ", " ")
        .replace("\n", "")
    )


def faiFile(df):
    fileName = df["Sede"].iat[0] + ".docx"
    df["testo"] = (
        df["titolo programma "]
        + "\n"
        + df["titolo progetto"]
        + " - "
        + df["x"].astype(int).astype(str)
        + " posti ordinari"
    )

    df["testo"] = df.apply(
        lambda row: (
            row["testo"] + " e " + str(int(row["y"])) + " posti GMO"
            if row["y"] > 0
            else row["testo"]
        ),
        axis=1,
    )

    testo_completo = "\n\n".join(df["testo"].astype(str))
    create_file(fileName, testo_completo)


# una testo word con un testo replacement
def replace(new_document, word, replacement):
    # print("replace", word, replacement)
    for p in new_document.paragraphs:
        if p.text.find(word) >= 0:
            p.text = p.text.replace(word, replacement)

    for table in new_document.tables:
        for r in table.rows:
            for c in r.cells:
                if c.text.find(word) >= 0:
                    c.text = c.text.replace(word, replacement)


# Crea il file "sede_nome_cognome.docx" riempiendo gli spazi indicati
def create_file(filename, text):
    new_document = Document(template_file)
    replace(new_document, "$$$", text)
    new_document.save(output_dir + filename)


## main

#     filename = filename + ".docx"

df = pd.read_excel(input_file, 0)
# df.iloc[:, 0] = df.iloc[:, 0].str.strip()
df["Sede"] = df.iloc[:, 0].apply(clean)
groups = df.groupby(df.columns[0])[df.columns].apply(faiFile)

if len(check):
    print("Possibili ERRORI:")
    print("\n".join(check))
